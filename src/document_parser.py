import PyPDF2
import pdfplumber
from docx import Document
import os
import re


class DocumentParser:
    """Parse documents and extract text content"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_file(self, file_path):
        """
        Parse a file and extract text based on its format
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text content
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.parse_pdf(file_path)
        elif file_extension == '.docx':
            return self.parse_docx(file_path)
        elif file_extension == '.txt':
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def parse_pdf(self, file_path):
        """
        Extract text from PDF file
        Uses pdfplumber as primary method, falls back to PyPDF2
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
        """
        text = ""
        
        try:
            # Method 1: pdfplumber (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber fails, try PyPDF2
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
        
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        
        return self.clean_extracted_text(text)
    
    def parse_docx(self, file_path):
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return self.clean_extracted_text(text)
        
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def parse_txt(self, file_path):
        """
        Extract text from TXT file
        
        Args:
            file_path (str): Path to TXT file
            
        Returns:
            str: Extracted text
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return self.clean_extracted_text(text)
        
        except UnicodeDecodeError:
            # Try different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return self.clean_extracted_text(text)
            except Exception as e:
                raise Exception(f"Error parsing TXT: {str(e)}")
        
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    def clean_extracted_text(self, text):
        """
        Clean extracted text by removing extra whitespace and special characters
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple newlines
        text = re.sub(r'\n+', '\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def is_supported(self, filename):
        """
        Check if file format is supported
        
        Args:
            filename (str): Name of the file
            
        Returns:
            bool: True if supported, False otherwise
        """
        file_extension = os.path.splitext(filename)[1].lower()
        return file_extension in self.supported_formats
    
    def extract_sections(self, text):
        """
        Extract common CV/Job sections using pattern matching
        
        Args:
            text (str): Full document text
            
        Returns:
            dict: Dictionary with extracted sections
        """
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'full_text': text
        }
        
        # Common section headers (case-insensitive)
        skills_patterns = [
            r'(?i)(skills?|technical skills?|competencies)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)',
            r'(?i)(expertise|proficiencies)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)'
        ]
        
        experience_patterns = [
            r'(?i)(experience|work experience|employment history)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)',
            r'(?i)(professional experience|career history)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)'
        ]
        
        education_patterns = [
            r'(?i)(education|academic background|qualifications?)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)',
            r'(?i)(degrees?|certifications?)[\s:]+([^\n]+(?:\n(?!\n)[^\n]+)*)'
        ]
        
        # Extract skills
        for pattern in skills_patterns:
            match = re.search(pattern, text)
            if match:
                sections['skills'] = match.group(2).strip()
                break
        
        # Extract experience
        for pattern in experience_patterns:
            match = re.search(pattern, text)
            if match:
                sections['experience'] = match.group(2).strip()
                break
        
        # Extract education
        for pattern in education_patterns:
            match = re.search(pattern, text)
            if match:
                sections['education'] = match.group(2).strip()
                break
        
        return sections


# Example usage
if __name__ == "__main__":
    parser = DocumentParser()
    
    # Test PDF parsing
    try:
        text = parser.parse_file("sample.pdf")
        print("PDF Text:", text[:500])
    except Exception as e:
        print("PDF Error:", e)
    
    # Test DOCX parsing
    try:
        text = parser.parse_file("sample.docx")
        print("DOCX Text:", text[:500])
    except Exception as e:
        print("DOCX Error:", e)