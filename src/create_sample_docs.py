"""
Create sample CV and Job Description documents for testing
"""

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def create_sample_cv_docx():
    """Create a sample CV in DOCX format"""
    doc = Document()
    
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Senior Data Scientist')
    
    doc.add_heading('SKILLS', 1)
    doc.add_paragraph(
        'Python, Machine Learning, TensorFlow, Keras, PyTorch, Scikit-learn, '
        'Pandas, NumPy, Data Analysis, Deep Learning, Neural Networks, '
        'Computer Vision, NLP, SQL, Big Data, Spark, Docker, Git'
    )
    
    doc.add_heading('EXPERIENCE', 1)
    doc.add_paragraph('Senior Data Scientist at Tech Corp (2020-2024)')
    doc.add_paragraph(
        '• Led a team of 5 data scientists in developing ML models for customer segmentation\n'
        '• Implemented recommendation systems that increased revenue by 30%\n'
        '• Built and deployed 10+ production ML models using TensorFlow and PyTorch\n'
        '• Designed data pipelines processing 1M+ records daily using Apache Spark\n'
        '• Mentored junior data scientists and conducted code reviews'
    )
    
    doc.add_paragraph('Data Scientist at AI Solutions (2018-2020)')
    doc.add_paragraph(
        '• Developed predictive models for financial forecasting\n'
        '• Created NLP solutions for text classification and sentiment analysis\n'
        '• Collaborated with engineering teams to deploy ML models to production\n'
        '• Improved model accuracy by 25% through feature engineering'
    )
    
    doc.add_heading('EDUCATION', 1)
    doc.add_paragraph('Master of Science in Computer Science')
    doc.add_paragraph('Stanford University (2016-2018)')
    doc.add_paragraph('Specialization: Machine Learning and Artificial Intelligence')
    
    doc.add_paragraph('Bachelor of Science in Mathematics')
    doc.add_paragraph('MIT (2012-2016)')
    
    doc.add_heading('CERTIFICATIONS', 1)
    doc.add_paragraph('• TensorFlow Developer Certificate')
    doc.add_paragraph('• AWS Certified Machine Learning - Specialty')
    doc.add_paragraph('• Deep Learning Specialization (Coursera)')
    
    doc.save('sample_cv.docx')
    print("✓ Created sample_cv.docx")

def create_sample_job_docx():
    """Create a sample job description in DOCX format"""
    doc = Document()
    
    doc.add_heading('Senior Data Scientist', 0)
    doc.add_heading('Job Description', 1)
    
    doc.add_paragraph(
        'We are seeking an experienced Data Scientist to join our AI team. '
        'The ideal candidate will have strong expertise in machine learning, '
        'deep learning, and production ML systems.'
    )
    
    doc.add_heading('REQUIRED SKILLS', 1)
    doc.add_paragraph(
        'Python, Machine Learning, TensorFlow or PyTorch, Scikit-learn, '
        'Deep Learning, Neural Networks, NLP, Computer Vision, '
        'SQL, Big Data Technologies (Spark/Hadoop), '
        'Data Visualization, Statistical Analysis, Git'
    )
    
    doc.add_heading('EXPERIENCE REQUIRED', 1)
    doc.add_paragraph(
        '• 5+ years of experience in data science or machine learning roles\n'
        '• Proven track record of deploying ML models to production\n'
        '• Experience leading data science projects and teams\n'
        '• Strong background in statistical modeling and analysis\n'
        '• Experience with cloud platforms (AWS/GCP/Azure)\n'
        '• Excellent problem-solving and communication skills'
    )
    
    doc.add_heading('EDUCATION REQUIRED', 1)
    doc.add_paragraph(
        "Master's degree or PhD in Computer Science, Statistics, Mathematics, "
        "or related field. Bachelor's degree with exceptional experience may be considered."
    )
    
    doc.add_heading('RESPONSIBILITIES', 1)
    doc.add_paragraph(
        '• Design and implement machine learning solutions for business problems\n'
        '• Build and optimize deep learning models\n'
        '• Collaborate with engineering teams for model deployment\n'
        '• Mentor junior data scientists\n'
        '• Present findings to stakeholders\n'
        '• Stay current with latest ML research and technologies'
    )
    
    doc.add_heading('BENEFITS', 1)
    doc.add_paragraph('• Competitive salary and equity')
    doc.add_paragraph('• Health, dental, and vision insurance')
    doc.add_paragraph('• Flexible work arrangements')
    doc.add_paragraph('• Professional development budget')
    doc.add_paragraph('• Collaborative team environment')
    
    doc.save('sample_job.docx')
    print("✓ Created sample_job.docx")

def create_sample_cv_txt():
    """Create a sample CV in TXT format"""
    content = """
JOHN DOE
Senior Data Scientist
Email: john.doe@email.com | Phone: (555) 123-4567

SKILLS
Python, Machine Learning, TensorFlow, Keras, PyTorch, Scikit-learn, Pandas, NumPy, 
Data Analysis, Deep Learning, Neural Networks, Computer Vision, NLP, SQL, Big Data, 
Spark, Docker, Git, AWS, Statistical Analysis, Data Visualization

EXPERIENCE

Senior Data Scientist | Tech Corp | 2020-2024
• Led a team of 5 data scientists in developing ML models for customer segmentation
• Implemented recommendation systems that increased revenue by 30%
• Built and deployed 10+ production ML models using TensorFlow and PyTorch
• Designed data pipelines processing 1M+ records daily using Apache Spark
• Mentored junior data scientists and conducted code reviews

Data Scientist | AI Solutions | 2018-2020
• Developed predictive models for financial forecasting
• Created NLP solutions for text classification and sentiment analysis
• Collaborated with engineering teams to deploy ML models to production
• Improved model accuracy by 25% through feature engineering

EDUCATION

Master of Science in Computer Science
Stanford University | 2016-2018
Specialization: Machine Learning and Artificial Intelligence

Bachelor of Science in Mathematics
MIT | 2012-2016

CERTIFICATIONS
• TensorFlow Developer Certificate
• AWS Certified Machine Learning - Specialty
• Deep Learning Specialization (Coursera)
"""
    
    with open('sample_cv.txt', 'w') as f:
        f.write(content)
    print("✓ Created sample_cv.txt")

def create_sample_job_txt():
    """Create a sample job description in TXT format"""
    content = """
SENIOR DATA SCIENTIST
Job ID: J001

JOB DESCRIPTION
We are seeking an experienced Data Scientist to join our AI team. The ideal candidate 
will have strong expertise in machine learning, deep learning, and production ML systems.

REQUIRED SKILLS
Python, Machine Learning, TensorFlow or PyTorch, Scikit-learn, Deep Learning, 
Neural Networks, NLP, Computer Vision, SQL, Big Data Technologies (Spark/Hadoop), 
Data Visualization, Statistical Analysis, Git

EXPERIENCE REQUIRED
• 5+ years of experience in data science or machine learning roles
• Proven track record of deploying ML models to production
• Experience leading data science projects and teams
• Strong background in statistical modeling and analysis
• Experience with cloud platforms (AWS/GCP/Azure)
• Excellent problem-solving and communication skills

EDUCATION REQUIRED
Master's degree or PhD in Computer Science, Statistics, Mathematics, or related field. 
Bachelor's degree with exceptional experience may be considered.

RESPONSIBILITIES
• Design and implement machine learning solutions for business problems
• Build and optimize deep learning models
• Collaborate with engineering teams for model deployment
• Mentor junior data scientists
• Present findings to stakeholders
• Stay current with latest ML research and technologies

BENEFITS
• Competitive salary and equity
• Health, dental, and vision insurance
• Flexible work arrangements
• Professional development budget
• Collaborative team environment
"""
    
    with open('sample_job.txt', 'w') as f:
        f.write(content)
    print("✓ Created sample_job.txt")

if __name__ == "__main__":
    print("\n" + "="*50)
    print("Creating Sample Documents")
    print("="*50 + "\n")
    
    try:
        # Create DOCX files
        create_sample_cv_docx()
        create_sample_job_docx()
        
        # Create TXT files
        create_sample_cv_txt()
        create_sample_job_txt()
        
        print("\n" + "="*50)
        print("✓ All sample documents created successfully!")
        print("="*50)
        print("\nYou can now use these files to test the upload functionality:")
        print("• sample_cv.docx")
        print("• sample_cv.txt")
        print("• sample_job.docx")
        print("• sample_job.txt")
        
    except Exception as e:
        print(f"\n✗ Error creating documents: {str(e)}")
        print("\nMake sure you have installed python-docx:")
        print("pip install python-docx")